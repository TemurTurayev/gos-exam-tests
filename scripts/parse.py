#!/usr/bin/env python3
"""Parse the source test files (with known answer keys) into unified JSON.

Каждый исходник размечает правильный ответ по-своему, поэтому здесь несколько
парсеров. Нумерованные форматы (C/D/E/F) используют общий алгоритм
`parse_numbered`, который устойчив к «грязной» разметке:

* номер вопроса может быть без точки (`23   Текст`);
* метка варианта может быть `А.`, `А)` или `А   ` (буква + пробелы);
* строки-продолжения приклеиваются к предыдущему варианту, а не создают новый;
* новый вопрос принимается только если его номер продолжает нумерацию —
  это защищает от строк вроде «30-32», которые иначе выглядят как «вопрос 30».
"""
import json
import re
import subprocess
from pathlib import Path

import docx
import fitz  # PyMuPDF

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "sources" / "with_answers"
DATA = ROOT / "data"
DATA.mkdir(exist_ok=True)

CYR_LAT = "A-Za-zА-Яа-яЁёЎўҚқҒғҲҳ"

# "1. текст", "23   текст", "*5) текст"
Q_START = re.compile(r"^\s*[*#]?\s*(\d{1,4})\s*[.)]?\s*(.*)$")
# "А. текст", "#б) текст", "*А) текст", "Г    текст"
OPT_LABEL = re.compile(rf"^\s*([#*+]?)\s*([{CYR_LAT}])\s*(?:[.)]\s*|\s{{2,}})(.*)$")
# "B текст" — метка заглавной буквой через один пробел, без точки.
# Ограничиваем набором реальных меток вариантов, иначе под шаблон попадёт
# обычный текст, начинающийся с заглавной буквы.
OPT_LABEL_LOOSE = re.compile(r"^\s*([#*+]?)\s*([ABCDEFGАБВГДЕЖ])\s(\S.*)$")
# лишние маркеры в начале строки
LEAD_MARK = re.compile(r"^\s*[#*+\-]+\s*")


def strip_marks(text):
    t = LEAD_MARK.sub("", text).strip()
    t = t.rstrip("*").strip()
    return t


def textutil_txt(path: Path) -> str:
    out = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True, text=True,
    )
    return out.stdout


def lines_of(text):
    return [l for l in (x.strip() for x in text.splitlines()) if l]


def bold_ratio(paragraph):
    """Доля жирного текста в абзаце — так надёжнее, чем any(run.bold)."""
    total = bold = 0
    for r in paragraph.runs:
        n = len(r.text.strip())
        if not n:
            continue
        total += n
        if r.bold:
            bold += n
    return (bold / total) if total else 0.0


# ---------------------------------------------------------------- generic ----
BULLET = re.compile(r"^\s*[••·]\s*")


def parse_numbered(items, is_correct, gap=6, opt_mode="labeled", bullets=False):
    """items: список (text, meta). is_correct(raw_text, meta) -> bool.

    opt_mode="labeled" — варианты помечены буквой (`А.`, `б)`, `Г   `);
    opt_mode="any"     — вариантом считается любая строка, не начавшая вопрос.
    """
    questions = []
    cur = None
    last_num = 0

    def flush():
        nonlocal cur
        if cur and cur["options"]:
            questions.append(cur)
        cur = None

    def start(text):
        nonlocal cur
        flush()
        cur = {"q": text.strip(), "options": [], "correct": []}

    def add_option(raw, body, meta):
        if is_correct(raw, meta):
            cur["correct"].append(len(cur["options"]))
        cur["options"].append(strip_marks(body))

    for raw, meta in items:
        text = raw.strip()
        if not text or text.startswith("@"):
            continue

        if bullets and BULLET.match(text):
            start(BULLET.sub("", text))
            continue

        # 1) помеченный вариант ответа — проверяем ДО номера вопроса,
        #    иначе строка вида "А 30-32" может сойти за начало вопроса
        om = None
        if opt_mode == "labeled":
            om = OPT_LABEL.match(text)
            # «B текст» принимаем только если у вопроса уже есть варианты —
            # так строка вопроса не будет принята за вариант
            if not om and cur is not None and cur["options"]:
                om = OPT_LABEL_LOOSE.match(text)
        if om and om.group(3).strip() and cur is not None:
            add_option(text, om.group(3), meta)
            continue

        # 2) начало нового вопроса
        m = Q_START.match(text)
        if m and m.group(2).strip():
            num = int(m.group(1))
            has_opts = bool(cur and len(cur["options"]) >= 2)
            # обычный шаг нумерации, небольшой пропуск, либо ресинхронизация
            # после большого пропуска (в исходниках часть вопросов вырезана)
            if (num == last_num + 1
                    or last_num < num < last_num + gap
                    or (opt_mode == "labeled" and has_opts and num > last_num)):
                start(m.group(2))
                last_num = num
                continue

        if cur is None:
            continue

        # 3) непомеченный вариант ответа
        if opt_mode == "any" and cur["q"]:
            add_option(text, text, meta)
            continue

        # 4) продолжение предыдущей строки
        if cur["options"]:
            cur["options"][-1] += " " + strip_marks(text)
        else:
            cur["q"] += " " + text

    flush()
    return questions


# ---------- Format A: "#question" / "+correct" / "-wrong" ----------
def parse_format_a(lines):
    questions = []
    cur = None
    for idx, l in enumerate(lines):
        # у части вопросов в исходнике потерян ведущий «#»: распознаём их по
        # тому, что строка без маркера идёт перед строкой варианта «+»/«-»
        if (not l.startswith(("#", "+", "-"))
                and cur is not None and cur["options"]
                and idx + 1 < len(lines) and lines[idx + 1].startswith(("+", "-"))):
            questions.append(cur)
            cur = {"q": l.strip(), "options": [], "correct": []}
            continue
        if l.startswith("#"):
            if cur and cur["options"]:
                questions.append(cur)
            cur = {"q": l[1:].strip(), "options": [], "correct": []}
        elif l.startswith("+"):
            if cur is None:
                continue
            cur["correct"].append(len(cur["options"]))
            cur["options"].append(l[1:].strip().rstrip("*").strip())
        elif l.startswith("-"):
            if cur is None:
                continue
            cur["options"].append(l[1:].strip().rstrip("*").strip())
        else:
            if cur is None:
                continue
            if cur["options"]:
                cur["options"][-1] += " " + l.strip()
            else:
                cur["q"] += " " + l
    if cur and cur["options"]:
        questions.append(cur)
    return questions


# ---------- Format B: docx-таблица [сложность, вопрос, верный, неверные...] ----
def parse_format_b(path):
    d = docx.Document(str(path))
    questions = []
    for t in d.tables:
        header_skipped = False
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells]
            if not header_skipped:
                header_skipped = True
                if len(cells) > 1 and "тест" in cells[1].lower():
                    continue
            if len(cells) < 3 or not cells[1] or not cells[2]:
                continue
            options = [cells[2]] + [c for c in cells[3:] if c]
            questions.append({"q": cells[1], "options": options, "correct": [0]})
    return questions


# ---------- Format G: PDF "Правильный ответ: X. текст" ----------
def parse_format_g(path):
    doc = fitz.open(str(path))
    lines = lines_of("\n".join(page.get_text() for page in doc))
    questions = []
    cur = None
    correct_letter = correct_text = None

    def close():
        nonlocal cur
        if cur and cur["options"]:
            questions.append(cur)
        cur = None

    for l in lines:
        m = re.match(r"^Правильный ответ:\s*([A-ZА-ЯЁ])\.?\s*(.*)$", l)
        if m:
            correct_letter, correct_text = m.group(1), m.group(2).strip()
            # в этом PDF ключ идёт ПОСЛЕ строки вопроса, а не перед ней
            if cur is not None and cur["_letter"] is None:
                cur["_letter"], cur["_text"] = correct_letter, correct_text
            continue
        qm = re.match(r"^(\d{1,4})\.\s+(.*)$", l)
        if qm:
            close()
            # ключ в этом PDF идёт строкой ПОСЛЕ вопроса, поэтому новый вопрос
            # всегда стартует без буквы — иначе он унаследует ответ соседа
            cur = {"q": qm.group(2).strip(), "options": [], "correct": [],
                   "_letter": None, "_text": None}
            correct_letter = correct_text = None
            continue
        if cur is None:
            continue
        om = re.match(rf"^\s*([{CYR_LAT}])\s*[.)]\s*(.*)$", l)
        if om and om.group(2).strip():
            letter, body = om.group(1), om.group(2).strip()
            # второй вид ключа в этом PDF: «... ←правильный ответ» в строке варианта
            if "←" in body:
                cur["correct"].append(len(cur["options"]))
                body = body.split("←")[0].strip()
            elif cur["_letter"] and letter == cur["_letter"]:
                cur["correct"].append(len(cur["options"]))
            cur["options"].append(body)
        elif cur["options"]:
            cur["options"][-1] += " " + l.strip()
        else:
            cur["q"] += " " + l.strip()
    close()
    for q in questions:
        q.pop("_letter", None)
        q.pop("_text", None)
    return questions


def finalize(questions, min_options=2, max_options=12):
    """Отбрасывает мусор разбора и перенумеровывает ответы после чистки вариантов.

    `correct` — список индексов: часть вопросов («перечислите 5 форм…»)
    имеет несколько правильных ответов.
    """
    out = []
    for q in questions:
        correct = q["correct"]
        if not correct or not q["q"].strip():
            continue
        options = q["options"]
        if max(correct) >= len(options):
            continue

        kept = [(i, o.strip()) for i, o in enumerate(options) if o.strip()]
        old_to_new = {old: new for new, (old, _) in enumerate(kept)}
        new_correct = sorted({old_to_new[i] for i in correct if i in old_to_new})
        if not new_correct:
            continue

        opts = [" ".join(o.split()) for _, o in kept]
        if not (min_options <= len(opts) <= max_options):
            continue
        if len(set(opts)) != len(opts):          # дубли вариантов = мусор разбора
            continue
        if len(new_correct) >= len(opts):        # «правильно всё» — тоже мусор
            continue

        out.append({"q": " ".join(q["q"].split()),
                    "options": opts,
                    "correct": new_correct})
    return out



# ---------------------------------------------------- проверка ключей ----
# Отдельная от парсера логика: заново собирает из исходника множество
# текстов, помеченных как правильные. Используется и при сборке (фильтр),
# и в verify.py (отчёт).
LABEL = re.compile(rf"^\s*[#*+\-]*\s*(?:[{CYR_LAT}]\s*[.)]|[{CYR_LAT}]\s{{2,}})\s*")


def norm(s):
    """Нормализуем текст для сравнения: регистр, пунктуация, пробелы."""
    s = LABEL.sub("", s.strip())
    s = re.sub(r"[^\w\s]", " ", s.lower())
    return " ".join(s.split())


def marked_texts(spec):
    """Множество нормализованных текстов, помеченных в исходнике как верные."""
    path = SRC / spec["file"]
    fmt = spec["fmt"]
    marked = set()

    if fmt in ("a", "a_doc", "a_docx"):
        if fmt == "a_docx":
            lines = [p.text.strip() for p in docx.Document(str(path)).paragraphs if p.text.strip()]
        elif path.suffix == ".txt":
            lines = lines_of(path.read_text(encoding="utf-8-sig"))
        else:
            lines = lines_of(textutil_txt(path))
        for l in lines:
            if l.startswith("+"):
                marked.add(norm(l[1:].rstrip("*")))

    elif fmt == "b":
        for t in docx.Document(str(path)).tables:
            for r in t.rows:
                cells = [c.text.strip() for c in r.cells]
                if len(cells) >= 3 and cells[2]:
                    marked.add(norm(cells[2]))

    elif fmt == "num_hash":
        for l in lines_of(textutil_txt(path)):
            if l.lstrip().startswith("#"):
                marked.add(norm(l))

    elif fmt in ("num_star", "num_star_plain"):
        if path.suffix == ".docx":
            lines = [p.text.strip() for p in docx.Document(str(path)).paragraphs if p.text.strip()]
        else:
            lines = lines_of(textutil_txt(path))
        for l in lines:
            s = l.strip()
            if s.startswith("*") or s.rstrip().endswith("*"):
                marked.add(norm(s.strip("*")))

    elif fmt == "num_bold":
        p = path if path.suffix == ".docx" else convert_doc_to_docx(path, ROOT / "scripts" / "_docx_conv")
        for para in docx.Document(str(p)).paragraphs:
            if para.text.strip() and bold_ratio(para) > 0.5:
                marked.add(norm(para.text))

    elif fmt == "g":
        doc = fitz.open(str(path))
        lines = lines_of("\n".join(pg.get_text() for pg in doc))
        pending = None
        for l in lines:
            m = re.match(r"^Правильный ответ:\s*([A-ZА-ЯЁ])\.?\s*(.*)$", l)
            if m:
                pending = m.group(1)
                if m.group(2).strip():
                    marked.add(norm(m.group(2)))
                continue
            if "←" in l:
                marked.add(norm(l.split("←")[0]))
                continue
            om = re.match(rf"^\s*([{CYR_LAT}])\s*[.)]\s*(.*)$", l)
            if om and pending and om.group(1) == pending:
                marked.add(norm(om.group(2)))
                pending = None
    return marked


def convert_doc_to_docx(doc_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / (doc_path.stem + ".docx")
    subprocess.run(
        ["textutil", "-convert", "docx", "-output", str(target), str(doc_path)],
        capture_output=True, text=True,
    )
    return target


def items_from_lines(lines):
    return [(l, None) for l in lines]


def items_from_docx(path):
    d = docx.Document(str(path))
    return [(p.text, bold_ratio(p)) for p in d.paragraphs if p.text.strip()]


# детекторы правильного ответа для нумерованных форматов
def correct_by_hash(raw, meta):
    return raw.lstrip().startswith("#")


def correct_by_star(raw, meta):
    s = raw.strip()
    return s.startswith("*") or s.rstrip().endswith("*")


def correct_by_bold(raw, meta):
    return (meta or 0) > 0.5


FILES = [
    dict(id="ped-ambulator-ru", title="Педиатрия — амбулаторно-поликлиническая",
         subject="Педиатрия", language="ru", fmt="a",
         file="Ambulator-poliklinik pediatriya RUS.txt"),
    dict(id="ped-1000-uz", title="Педиатрия — 1000 тестов",
         subject="Педиатрия", language="uz", fmt="a",
         file="pediatriya 1000 full.txt"),
    dict(id="child-surg-hospital-uz", title="Госпитальная детская хирургия",
         subject="Детская хирургия", language="uz", fmt="num_hash",
         file="Госпитал болалар хирургияси.doc"),
    dict(id="child-surg-ru", title="Детская хирургия",
         subject="Детская хирургия", language="ru", fmt="num_bold",
         file="Детская хирургия.doc"),
    dict(id="therapy-internal-ru", title="Внутренние болезни",
         subject="Терапия", language="ru", fmt="num_hash",
         file="Ички касалликлар (Рус).doc"),
    dict(id="ped-1000-2-uz", title="Педиатрия — расширенный банк",
         subject="Педиатрия", language="uz", fmt="num_star",
         file="Педиатрия-1000.docx"),
    dict(id="therapy-1000-uz", title="Терапия — 1000 тестов",
         subject="Терапия", language="uz", fmt="num_bold",
         file="Терапия - 1000.doc"),
    dict(id="therapy-ecg-ru", title="Терапия — ЭКГ и кардиология",
         subject="Терапия", language="ru", fmt="g",
         file="Терапия ....✓.pdf"),
    dict(id="therapy-table-ru", title="Терапия — тесты по темам",
         subject="Терапия", language="ru", fmt="b",
         file="Терапия русс с ответами.docx"),
    dict(id="therapy-test-ru", title="Терапия — факультетская",
         subject="Терапия", language="ru", fmt="a_doc",
         file="Терапия тест.doc"),
    dict(id="surg-facult-ru", title="Факультетская хирургия",
         subject="Хирургия", language="ru", fmt="num_star",
         file="Фак. хирургия.doc"),
    dict(id="child-surg-facult-uz", title="Факультетская детская хирургия",
         subject="Детская хирургия", language="uz", fmt="num_hash",
         file="Факультет_болалар_хирургияси_Рус.doc"),
    dict(id="surg-hospital-ru", title="Госпитальная хирургия",
         subject="Хирургия", language="ru", fmt="a_docx",
         file="Хирургия тест.docx"),
    dict(id="obgyn-ru", title="Акушерство и гинекология",
         subject="Акушерство и гинекология", language="ru", fmt="num_star_plain",
         file="акушерлик ва гинекология рус.doc"),
    dict(id="ped-facult-ru", title="Факультетская педиатрия",
         subject="Педиатрия", language="ru", fmt="num_star",
         file="педиатрия  рус.doc"),
    dict(id="surg-1000-uz", title="Хирургия — 1000 тестов",
         subject="Хирургия", language="uz", fmt="num_bold",
         file="хирургия - 1000.doc"),
]


def parse_one(spec, conv_dir):
    path = SRC / spec["file"]
    fmt = spec["fmt"]

    if fmt == "a":
        text = path.read_text(encoding="utf-8-sig") if path.suffix == ".txt" else textutil_txt(path)
        return parse_format_a(lines_of(text))
    if fmt == "a_doc":
        return parse_format_a(lines_of(textutil_txt(path)))
    if fmt == "a_docx":
        return parse_format_a([t.strip() for t, _ in items_from_docx(path) if t.strip()])
    if fmt == "b":
        return parse_format_b(path)
    if fmt == "g":
        return parse_format_g(path)
    if fmt == "num_hash":
        return parse_numbered(items_from_lines(lines_of(textutil_txt(path))), correct_by_hash)
    if fmt == "num_star":
        if path.suffix == ".docx":
            items = [(t, None) for t, _ in items_from_docx(path)]
        else:
            items = items_from_lines(lines_of(textutil_txt(path)))
        return parse_numbered(items, correct_by_star, bullets=True)
    if fmt == "num_star_plain":
        items = items_from_lines(lines_of(textutil_txt(path)))
        return parse_numbered(items, correct_by_star, opt_mode="any")
    if fmt == "num_bold":
        docx_path = path if path.suffix == ".docx" else convert_doc_to_docx(path, conv_dir)
        return parse_numbered(items_from_docx(docx_path), correct_by_bold)
    raise ValueError(fmt)


def apply_tags(questions, marked):
    """Помечает вопросы: verified — ответ подтверждён исходником, disputed — нет.

    Разметка в файле есть у всех этих вопросов, но независимая проверка
    сходится не всегда; такие вопросы не выбрасываем, а показываем с пометкой.
    """
    out = []
    for q in questions:
        ok = all(norm(q["options"][i]) in marked for i in q["correct"])
        out.append({**q, "tag": "verified" if ok else "disputed"})
    return out


def build_core():
    """Наборы из файлов, где правильный ответ размечен в самом документе."""
    conv_dir = ROOT / "scripts" / "_docx_conv"
    sets = {}
    for spec in FILES:
        questions = apply_tags(finalize(parse_one(spec, conv_dir)), marked_texts(spec))
        sets[spec["id"]] = {"id": spec["id"], "title": spec["title"],
                            "subject": spec["subject"], "language": spec["language"],
                            "questions": questions}
    return sets


def main():
    for sid, data in build_core().items():
        tags = {}
        for q in data["questions"]:
            tags[q["tag"]] = tags.get(q["tag"], 0) + 1
        print(f"{sid:24s} {len(data['questions']):5d}  {tags}")


if __name__ == "__main__":
    main()
